from PIL import Image
from docx import *
from docx.shared import *

class asciimage:
    """
    Instantiate an ascii image object.
    
    :param image_path: The path of the input image. Can provide absolute or relative path
    :type image_path: string
    :param file_path: Output image path: defines where the ascii image will be saved 
    :type file_path: string
    :param ascii_char: A list of characters used to draw the ascii image. By default the character list is "$@B%8&WM#*oahkbdpqwmZO0QLCJUYXzcvunxrjft/\|()1{}[]?-_+~<>i!lI;:,\"^`'. " 
    :type file_path: list
    :param height: desired height of output image 
    :type file_path: int
    :param width: desired width of output image 
    :type file_path: int
    """
    
    def __init__(self, image_path,file_path,ascii_char=list("$@B%8&WM#*oahkbdpqwmZO0QLCJUYXzcvunxrjft/\|()1{}[]?-_+~<>i!lI;:,\"^`'. "),height = 256,width = 256):
        self.image_path = image_path
        self.file_path = file_path
        self.ascii_char = ascii_char    
        self.height = height
        self.width = width
        

    def make_ascii_image(self):
        """
        produces an ASCII image in a text file 
        """
        from PIL import Image
        im = Image.open(self.image_path)#defined the image path
        im = im.resize((self.height, self.width), Image.NEAREST)#resizing using nearest interpolation
        txt = ""
        for i in range(self.height):
            for j in range(self.width):
                length = len(self.ascii_char)
                #checking image path jpg or png
                if self.image_path.split('.')[-1] == 'png':
                  (r,g,b,a) = im.getpixel((j, i))#getting rgb pixel values at image position(j,i), a for transparency
                else:
                  (r,g,b) = im.getpixel((j, i))#getting rgb pixel values at image position(j,i)  
                gray = int(0.2126 * r + 0.7152 * g + 0.0722 * b)#standard grayscale conversion
                unit = (256.0 + 1) / length
                txt += self.ascii_char[int(gray / unit)]#choosing a character from input ascii list
            txt += '\n'
        with open(self.file_path, 'w') as f:
            f.write(txt)
            
    def make_colored_ascii(self, size):
        """
        produces a colored ASCII image in a doc file 
        """
        im = Image.open(self.image_path)#defined the image path
        im = im.transpose(Image.ROTATE_90)
        im = im.resize((size, size))
        col = []
        for i in range(im.size[0]):
          colis = []
          for j in range(im.size[1]):
              if self.image_path.split('.')[-1] == 'png':
                  (r,g,b,a) = im.getpixel((i, j))
              else:
                (r,g,b) = im.getpixel((i, j))   
              
              colis.append([r,g,b])
          col.append(colis)
        doc = Document()
        doc_para = doc.add_paragraph('')
        doc_para.paragraph_format.line_spacing = Inches(0.025)
        for i in range(64):
          for j in range(64):
            doc_para.add_run("A").font.color.rgb = RGBColor(col[i][j][0],col[i][j][1],col[i][j][2])
          doc_para.add_run('\n')
        doc.save(self.file_path)        
            
